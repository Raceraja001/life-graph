"""File parsers for knowledge base ingestion.

Extracts text content from uploaded files for chunking and embedding.
Supported formats: .txt, .md, .csv, .pdf, .docx

PDF support requires pymupdf (fitz), DOCX requires python-docx.
Both degrade gracefully — if the library is missing, the parser
returns an error message instead of crashing.
"""

from __future__ import annotations

import csv
import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_text(content: bytes, filename: str) -> str:
    """Parse plain text and markdown files."""
    return content.decode("utf-8", errors="replace")


def parse_csv(content: bytes, filename: str) -> str:
    """Parse CSV files into row-per-line text."""
    text = content.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))

    lines: list[str] = []
    headers: list[str] = []

    for i, row in enumerate(reader):
        if i == 0:
            headers = row
            lines.append("Headers: " + " | ".join(headers))
            continue
        # Create readable key:value pairs per row
        if headers:
            pairs = [f"{h}: {v}" for h, v in zip(headers, row) if v.strip()]
            lines.append(f"Row {i}: " + ", ".join(pairs))
        else:
            lines.append(" | ".join(row))

    return "\n".join(lines)


def parse_pdf(content: bytes, filename: str) -> str:
    """Parse PDF files using pymupdf (fitz)."""
    try:
        import fitz  # pymupdf
    except ImportError:
        logger.warning("pymupdf not installed — cannot parse PDF files. Run: pip install pymupdf")
        return f"[Error: PDF parsing unavailable. Install pymupdf to parse {filename}]"

    doc = fitz.open(stream=content, filetype="pdf")
    pages: list[str] = []

    for i, page in enumerate(doc):
        text = page.get_text("text")
        if text.strip():
            pages.append(f"--- Page {i + 1} ---\n{text.strip()}")

    doc.close()
    return "\n\n".join(pages)


def parse_docx(content: bytes, filename: str) -> str:
    """Parse DOCX files using python-docx."""
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx not installed — cannot parse DOCX files. Run: pip install python-docx")
        return f"[Error: DOCX parsing unavailable. Install python-docx to parse {filename}]"

    doc = Document(io.BytesIO(content))
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n\n".join(paragraphs)


# ── File extension → parser mapping ───────────────────────

PARSERS: dict[str, callable] = {
    ".txt": parse_text,
    ".md": parse_text,
    ".markdown": parse_text,
    ".csv": parse_csv,
    ".pdf": parse_pdf,
    ".docx": parse_docx,
}

SUPPORTED_EXTENSIONS = set(PARSERS.keys())
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB


def parse_file(content: bytes, filename: str) -> str:
    """Parse a file based on its extension.

    Args:
        content: Raw file bytes
        filename: Original filename (used for extension detection)

    Returns:
        Extracted text content

    Raises:
        ValueError: If file type is not supported or file is too large
    """
    if len(content) > MAX_FILE_SIZE:
        raise ValueError(f"File too large: {len(content)} bytes (max {MAX_FILE_SIZE})")

    ext = Path(filename).suffix.lower()

    if ext not in PARSERS:
        raise ValueError(
            f"Unsupported file type: {ext}. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    parser = PARSERS[ext]
    text = parser(content, filename)

    if not text or not text.strip():
        raise ValueError(f"No text content extracted from {filename}")

    return text
